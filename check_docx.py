from docx import Document

# 必修1总结文件路径
docx_path = r'D:\BaiduNetdiskDownload\生物（2011-2019）\2011-2019上海生物合格考\生物合格考知识点总结【必修1】-2023年高中生物合格考得分秘籍（上海专用）.docx'

doc = Document(docx_path)
print(f'总段落数: {len(doc.paragraphs)}')
print('\n前15段内容:')
for i, p in enumerate(doc.paragraphs[:15]):
    if p.text.strip():
        print(f'{i+1}: {p.text.strip()[:100]}')